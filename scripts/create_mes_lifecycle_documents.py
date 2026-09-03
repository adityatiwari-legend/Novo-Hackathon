"""
Script to create authentic fictional "Novo Life MES PAS-X" dummy lifecycle documents
using python-docx in data/sample_documents/

These documents describe a fictional "Novo Life" MES PAS-X implementation.
They are explicitly DUMMY / HACKATHON / TRAINING SIMULATION records.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from backend.app.core.config import settings

def apply_dummy_watermark(doc, doc_id: str, title: str):
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = f"DUMMY / HACKATHON / TRAINING SIMULATION RECORD | {doc_id}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if hp.runs:
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(180, 50, 50)
        hp.runs[0].font.bold = True

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = "NOT A GENUINE NOVO NORDISK RECORD — STRICTLY FOR HACKATHON DEMO USE"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fp.runs:
        fp.runs[0].font.size = Pt(8)
        fp.runs[0].font.color.rgb = RGBColor(140, 140, 140)

def add_title_block(doc, doc_id: str, title: str, version: str, status: str, date_str: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sys = p.add_run("NOVO LIFE Computerized System Validation Package\n")
    r_sys.font.size = Pt(11)
    r_sys.font.bold = True
    r_sys.font.color.rgb = RGBColor(0, 89, 153)

    r_title = p.add_run(f"{title}\n")
    r_title.font.size = Pt(18)
    r_title.font.bold = True

    r_meta = p.add_run(f"Document ID: {doc_id}  |  Version: {version}  |  Status: {status}  |  Date: {date_str}\n")
    r_meta.font.size = Pt(9.5)
    r_meta.font.italic = True
    r_meta.font.color.rgb = RGBColor(90, 90, 90)

    # Simulation Notice Box
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box_table.cell(0, 0)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = cp.add_run("⚠️ DUMMY / HACKATHON / TRAINING SIMULATION RECORD\nThis document is a synthetic simulation record for the Novo Life MES PAS-X hackathon scenario. It does not represent genuine approved operational data.")
    c_run.font.size = Pt(8.5)
    c_run.font.bold = True
    c_run.font.color.rgb = RGBColor(160, 40, 40)
    doc.add_paragraph()

# -------------------------------------------------------------
# 1. NL-MES-MLGP-001: Master Lifecycle Generation Plan
# -------------------------------------------------------------
def build_mlgp():
    doc = Document()
    doc_id = "NL-MES-MLGP-001"
    title = "Master Lifecycle Generation Plan - Novo Life MES PAS-X"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "1.0", "Approved (Simulated)", "2025-01-15")

    doc.add_heading("1. Purpose & Scope", level=1)
    doc.add_paragraph(
        "This Master Lifecycle Generation Plan (MLGP) defines the computerized system validation "
        "and lifecycle management strategy for the Novo Life Manufacturing Execution System (MES) "
        "based on Werum PAS-X software (System ID: SYS-MES-001). This system is classified under GAMP 5 "
        "Category 4 (Configured Computerized System) and operates within a GxP-Critical commercial packaging environment."
    )

    doc.add_heading("2. System Classification & Governance", level=1)
    doc.add_paragraph(
        "• System Name: Novo Life MES PAS-X\n"
        "• System Identifier: SYS-MES-001\n"
        "• GAMP 5 Category: Category 4 (Configured Software)\n"
        "• GxP Impact: Direct Impact on Product Quality, Batch Release, and 21 CFR Part 11 Electronic Records\n"
        "• System Owner: Sarah Jenkins (IT Quality & Validated Systems)\n"
        "• Quality Unit Reviewer: Dr. Elena Rostova (QA Compliance)"
    )

    doc.add_heading("3. Lifecycle Phase Gates", level=1)
    doc.add_paragraph(
        "The project enforces six formal phase gates. Advancement across gates requires explicit Quality Unit authorization:\n"
        "• Gate G1: Project Inception & Regulatory Categorization\n"
        "• Gate G2: Requirements Baseline Approval (URS signed)\n"
        "• Gate G3: Design & Specification Freeze (FS, CS signed)\n"
        "• Gate G4: Verification & Qualification (IQ, Integration Testing, OV/PfV/UAT)\n"
        "• Gate G5: Release Readiness & Residual Risk Acceptance (VSR approval)\n"
        "• Gate G6: Operational Handover & SLA Activation"
    )
    return doc

# -------------------------------------------------------------
# 2. NL-MES-URS-001: User Requirement Specification (50 URS)
# -------------------------------------------------------------
def build_urs():
    doc = Document()
    doc_id = "NL-MES-URS-001"
    title = "User Requirement Specification - Novo Life MES PAS-X"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "1.0", "Baseline Approved (Simulated)", "2025-02-01")

    doc.add_heading("1. Executive Summary & Objective", level=1)
    doc.add_paragraph(
        "This document defines the 50 baseline user requirements for the Novo Life MES PAS-X solution (SYS-MES-001). "
        "The baseline consists of exactly 25 Functional Requirements and 25 Non-Functional Requirements. "
        "All requirements must maintain bidirectional traceability to design specifications, risk assessment items, "
        "and qualification protocols."
    )

    doc.add_heading("2. Functional Requirements (25 Baseline URS)", level=1)
    functional_reqs = [
        ("URS-001", "Master Batch Record (MBR) creation, lifecycle versioning, and electronic routing."),
        ("URS-002", "Weighing and dispensing material verification with electronic tolerance interlocks."),
        ("URS-003", "Electronic Signatures adhering to 21 CFR Part 11 and EU Annex 11 dual-witness criteria."),
        ("URS-004", "Material genealogy and end-to-end container tracking from dispensary to primary packaging."),
        ("URS-005", "Equipment calibration, cleaning status, and maintenance lockout interlocks."),
        ("URS-006", "In-Process Control (IPC) real-time data capture with automated limit verification."),
        ("URS-007", "Deviation and exception logging with automated event capture during batch execution."),
        ("URS-008", "Review by Exception (RBE) batch reporting highlighting out-of-spec parameters."),
        ("URS-009", "ERP (SAP S/4HANA) bidirectional interface for Process Orders and material consumption."),
        ("URS-010", "LIMS bi-directional interface for analytical sample creation and result synchronization."),
        ("URS-011", "Warehouse Management (WMS) barcode verification and container release validation."),
        ("URS-012", "Recipe authoring conforming to ISA-88 batch control standards."),
        ("URS-013", "Critical Process Parameter (CPP) alarm bounds enforcement and operator override logging."),
        ("URS-014", "Granular Role-Based Access Control (RBAC) ensuring segregation of shopfloor duties."),
        ("URS-015", "Contemporaneous electronic audit trail recording all parameter edits and acknowledgments."),
        ("URS-016", "Bill of Materials (BOM) quantity reconciliation and automatic yield calculations."),
        ("URS-017", "Workstation handheld and fixed 2D barcode scanner hardware interfacing."),
        ("URS-018", "Precision benchtop weigh scale serial RS-232 and Ethernet TCP/IP connectivity."),
        ("URS-019", "Automated PDF batch dossier generation with digital document sealing."),
        ("URS-020", "Container cleanliness, sanitization expiration, and room classification tracking."),
        ("URS-021", "Operator training qualification gating via LMS API prior to critical step access."),
        ("URS-022", "Sampling plan auto-generation and barcoded sample label thermal printing."),
        ("URS-023", "Environmental monitoring (cleanroom particle counter/temp) automated data linkage."),
        ("URS-024", "Master recipe approval workflow state machine (Draft, Review, Approved, Obsolete)."),
        ("URS-025", "Batch execution pause, hold, resume, and controlled abort emergency controls.")
    ]
    t_f = doc.add_table(rows=1, cols=3)
    t_f.rows[0].cells[0].paragraphs[0].add_run("Req ID").bold = True
    t_f.rows[0].cells[1].paragraphs[0].add_run("Classification").bold = True
    t_f.rows[0].cells[2].paragraphs[0].add_run("Requirement Statement").bold = True
    for r_id, desc in functional_reqs:
        row = t_f.add_row()
        row.cells[0].paragraphs[0].text = r_id
        row.cells[1].paragraphs[0].text = "FUNCTIONAL"
        row.cells[2].paragraphs[0].text = desc

    doc.add_heading("3. Non-Functional Requirements (25 Baseline URS)", level=1)
    non_functional_reqs = [
        ("URS-026", "High Availability: System availability shall exceed 99.9% across active production shifts."),
        ("URS-027", "Performance: Peak transaction response time shall not exceed 1.5 seconds under full load."),
        ("URS-028", "Audit Trail Display: Electronic audit trail shall be viewable and filterable within the application."),
        ("URS-029", "Disaster Recovery: RTO < 4 hours and RPO < 15 minutes for shopfloor recovery."),
        ("URS-030", "Automated Hourly Database Backup: Backups shall execute without transaction downtime."),
        ("URS-031", "Active Directory / Azure SSO: Authentication shall integrate with corporate IdP."),
        ("URS-032", "Inactivity Timeout: Workstation sessions shall lock after 15 minutes of idle time."),
        ("URS-033", "Encryption in Transit: All internal and external network communication shall use TLS 1.3."),
        ("URS-034", "Encryption at Rest: Database volumes and exported records shall use AES-256 encryption."),
        ("URS-035", "Audit Trail Immutability: Audit records shall be append-only with SHA-256 verification."),
        ("URS-036", "Concurrency: Support at least 250 simultaneous active shopfloor packaging operators."),
        ("URS-037", "Browser Compatibility: Responsive web client shall support Chromium and Edge enterprise builds."),
        ("URS-038", "Network Resilience: Workstations shall tolerate network latency up to 150 milliseconds."),
        ("URS-039", "Operating System Support: Server instances shall run on validated RHEL 9 or Windows Server 2022."),
        ("URS-040", "Time Synchronization: Server clocks shall synchronize with NTP stratum 1 within 100ms."),
        ("URS-041", "Crash Recovery: System crash during step execution shall recover without data corruption."),
        ("URS-042", "Data Retention: All batch records and audit logs shall be retained for 25 years."),
        ("URS-043", "Multi-Language Support: Shopfloor prompts shall support both English and Danish."),
        ("URS-044", "Audit Trail Export: Export format shall include human-readable PDF and delimited CSV."),
        ("URS-045", "Modular Scalability: System architecture shall allow modular expansion to future packaging lines."),
        ("URS-046", "Antivirus Compatibility: Real-time endpoint agent protection shall not impede serial communication."),
        ("URS-047", "Health Telemetry: Provide automated REST API endpoint for system operational monitoring."),
        ("URS-048", "Maintenance Banner: Display scheduled downtime warning to operators 2 hours in advance."),
        ("URS-049", "Offline Workstation Cache: Allow buffered local barcode scanning during 30-second network blips."),
        ("URS-050", "ALCOA+ Data Integrity: Ensure records are Attributable, Legible, Contemporaneous, Original, and Accurate.")
    ]
    t_nf = doc.add_table(rows=1, cols=3)
    t_nf.rows[0].cells[0].paragraphs[0].add_run("Req ID").bold = True
    t_nf.rows[0].cells[1].paragraphs[0].add_run("Classification").bold = True
    t_nf.rows[0].cells[2].paragraphs[0].add_run("Requirement Statement").bold = True
    for r_id, desc in non_functional_reqs:
        row = t_nf.add_row()
        row.cells[0].paragraphs[0].text = r_id
        row.cells[1].paragraphs[0].text = "NON-FUNCTIONAL"
        row.cells[2].paragraphs[0].text = desc

    doc.add_heading("4. Document Approvals & Signatures", level=1)
    doc.add_paragraph(
        "• Business Owner: Sarah Jenkins — Approved 2025-02-01\n"
        "• IT Validation Lead: Henrik Lindqvist — Approved 2025-02-01\n"
        "• Quality Assurance (QA Compliance Unit): MISSING - PENDING RE-AUTHORIZATION (Simulated Gap)"
    )
    return doc

# -------------------------------------------------------------
# 3. NL-MES-FS-001: Functional Specification
# -------------------------------------------------------------
def build_fs():
    doc = Document()
    doc_id = "NL-MES-FS-001"
    title = "Functional Specification - Novo Life MES PAS-X Core"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "1.0", "Baseline Approved", "2025-02-20")

    doc.add_heading("1. Module Architecture", level=1)
    doc.add_paragraph(
        "This Functional Specification maps the 50 URS requirements to Werum PAS-X software modules: "
        "MBR Design, Production Execution, Material Tracking, Equipment Management, and System Administration. "
        "All functional specifications are linked to the system requirement baseline in NL-MES-URS-001."
    )
    return doc

# -------------------------------------------------------------
# 4. NL-MES-ITRA-001: IT Risk Assessment (26 System Risks)
# -------------------------------------------------------------
def build_itra():
    doc = Document()
    doc_id = "NL-MES-ITRA-001"
    title = "IT Risk Assessment - Novo Life MES PAS-X (ICH Q9)"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "1.0", "Baseline Evaluated", "2025-03-01")

    doc.add_heading("1. Methodology & Scope", level=1)
    doc.add_paragraph(
        "This assessment establishes the initial system-level risk baseline using ICH Q9 Quality Risk Management principles. "
        "Twenty-six system hazards (RSK-MES-001 through RSK-MES-026) were identified across software architecture, "
        "interfacing, shopfloor dispensing, electronic record integrity, and disaster recovery."
    )

    doc.add_heading("2. System Risk Register (RSK-MES-001 to RSK-MES-026)", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].paragraphs[0].add_run("Risk ID").bold = True
    table.rows[0].cells[1].paragraphs[0].add_run("Hazard Description").bold = True
    table.rows[0].cells[2].paragraphs[0].add_run("Impact Area").bold = True
    table.rows[0].cells[3].paragraphs[0].add_run("Baseline Severity").bold = True

    risk_samples = [
        ("RSK-MES-001", "Unauthorized batch recipe modification leading to out-of-spec product packaging.", "Patient Safety", "HIGH"),
        ("RSK-MES-002", "Weigh scale serial communication buffer overrun resulting in incorrect active ingredient dispensing.", "Product Quality", "HIGH"),
        ("RSK-MES-003", "Electronic signature token interception or bypass in 21 CFR Part 11 validation.", "Data Integrity", "HIGH"),
        ("RSK-MES-004", "Material lot barcode mismatch resulting in unreleased intermediate processing.", "Product Quality", "HIGH"),
        ("RSK-MES-005", "Uncalibrated equipment interlock failure allowing out-of-tolerance run.", "Product Quality", "HIGH"),
        ("RSK-MES-006", "IPC real-time limits corruption causing undetected critical parameter excursion.", "Patient Safety", "HIGH"),
        ("RSK-MES-007", "Audit trail record alteration or truncation during catastrophic database failover.", "Data Integrity", "HIGH"),
        ("RSK-MES-008", "SAP S/4HANA Process Order synchronization desync causing duplicate packaging lot.", "Operational", "HIGH"),
        ("RSK-MES-009", "LIMS analytical sample status desync allowing unreleased batch packaging.", "Patient Safety", "HIGH"),
        ("RSK-MES-010", "WMS container status bypass resulting in quarantined material usage.", "Product Quality", "HIGH"),
        ("RSK-MES-011", "ISA-88 recipe phase deadlock halting primary packaging line.", "Operational", "HIGH"),
        ("RSK-MES-012", "CPP alarm suppression masking sterile boundary temperature excursion.", "Patient Safety", "HIGH"),
        ("RSK-MES-013", "Privileged administrative role escalation bypassing segregation of duties.", "Security / GxP", "HIGH"),
        ("RSK-MES-014", "Audit trail timestamp clock drift exceeding NTP tolerance.", "Data Integrity", "HIGH"),
        ("RSK-MES-015", "Yield calculation rounding error concealing bulk tablet loss.", "Regulatory", "HIGH"),
        ("RSK-MES-016", "Handheld scanner symbology decode error scanning incorrect label.", "Product Quality", "HIGH"),
        ("RSK-MES-017", "Serial interface noise introducing corrupt tare weight value.", "Product Quality", "HIGH"),
        ("RSK-MES-018", "PDF batch record generation timeout preventing batch release.", "Operational", "HIGH"),
        ("RSK-MES-019", "Room cleanliness expiration bypass allowing packaging in uncleaned suite.", "Product Quality", "HIGH"),
        ("RSK-MES-020", "Unqualified operator performing critical sterilizer unload step.", "GxP Compliance", "HIGH"),
        ("RSK-MES-021", "Barcode printer ribbon degradation rendering sample label unreadable.", "Data Integrity", "HIGH"),
        ("RSK-MES-022", "Particle counter environmental data packet loss during filling run.", "Data Integrity", "HIGH"),
        ("RSK-MES-023", "Unapproved recipe draft routed to shopfloor execution.", "Product Quality", "HIGH"),
        ("RSK-MES-024", "Uncontrolled batch abort leaving intermediate in unvalidated state.", "Product Quality", "HIGH"),
        ("RSK-MES-025", "Disaster recovery replica lag exceeding 15-minute RPO window.", "Business Continuity", "HIGH"),
        ("RSK-MES-026", "Electronic audit trail viewer filter truncation obscuring historic edits.", "Data Integrity", "MEDIUM")
    ]
    for r_id, desc, impact, sev in risk_samples:
        r = table.add_row()
        r.cells[0].paragraphs[0].text = r_id
        r.cells[1].paragraphs[0].text = desc
        r.cells[2].paragraphs[0].text = impact
        r.cells[3].paragraphs[0].text = sev

    return doc

# -------------------------------------------------------------
# 5. NL-MES-ITRRA-001: IT Requirement Risk Assessment
# -------------------------------------------------------------
def build_itrra():
    doc = Document()
    doc_id = "NL-MES-ITRRA-001"
    title = "IT Requirement Risk Assessment (ITRRA) - Novo Life MES PAS-X"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "1.0", "Draft / In Review", "2025-03-10")

    doc.add_heading("1. Executive Summary & Baseline Statistics", level=1)
    doc.add_paragraph(
        "This ITRRA cross-references all 50 URS requirements from NL-MES-URS-001 against the 26 system risk items "
        "defined in NL-MES-ITRA-001.\n\n"
        "Working Risk Classification Breakdown:\n"
        "• Total Evaluated Requirements: 50\n"
        "• Working High Risk Requirements: 49\n"
        "• Working Medium Risk Requirements: 1 (URS-028: Audit Trail Display & Filtering)\n"
        "• Working Low Risk Requirements: 0"
    )

    doc.add_heading("2. Critical Audit Observation: Residual Risk Status", level=1)
    p_warn = doc.add_paragraph()
    r = p_warn.add_run(
        "CRITICAL REGULATORY NOTICE: Residual risk is NOT RATED. "
        "While working risks have been classified, mitigation verification has not been completed, "
        "and residual risks have not been accepted by the Quality Unit. All 49 working high risks remain OPEN."
    )
    r.bold = True
    r.font.color.rgb = RGBColor(180, 20, 20)

    doc.add_heading("3. Requirement Risk Mapping Excerpt", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].paragraphs[0].add_run("URS ID").bold = True
    table.rows[0].cells[1].paragraphs[0].add_run("System Risk ID").bold = True
    table.rows[0].cells[2].paragraphs[0].add_run("Working Risk").bold = True
    table.rows[0].cells[3].paragraphs[0].add_run("Residual Risk State").bold = True

    mappings = [
        ("URS-001", "RSK-MES-001", "HIGH", "NOT RATED / OPEN"),
        ("URS-002", "RSK-MES-002", "HIGH", "NOT RATED / OPEN"),
        ("URS-003", "RSK-MES-003", "HIGH", "NOT RATED / OPEN"),
        ("URS-007", "RSK-MES-007", "HIGH", "NOT RATED / OPEN"),
        ("URS-015", "RSK-MES-014", "HIGH", "NOT RATED / OPEN"),
        ("URS-028", "RSK-MES-026", "MEDIUM", "NOT RATED / OPEN"),
        ("URS-050", "RSK-MES-007", "HIGH", "NOT RATED / OPEN")
    ]
    for uid, rid, wrk, res in mappings:
        row = table.add_row()
        row.cells[0].paragraphs[0].text = uid
        row.cells[1].paragraphs[0].text = rid
        row.cells[2].paragraphs[0].text = wrk
        row.cells[3].paragraphs[0].text = res
    return doc

# -------------------------------------------------------------
# 6. NL-MES-SUPA-001: Supplier Assessment
# -------------------------------------------------------------
def build_supa():
    doc = Document()
    doc_id = "NL-MES-SUPA-001"
    title = "Supplier Quality Assessment - Werum IT Solutions / Körber"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "1.0", "Approved with Conditions", "2024-11-20")

    doc.add_heading("1. Supplier Qualification Summary", level=1)
    doc.add_paragraph(
        "Assessment of Werum IT Solutions (Körber Pharma Software) for PAS-X MES v3.3. "
        "The vendor maintains an ISO 9001 certified Quality Management System. "
        "Overall supplier status is: QUALIFIED WITH CONDITIONS.\n\n"
        "Open Conditions:\n"
        "1. Periodic audit response pending for automated regression suite coverage.\n"
        "2. Formal escrow agreement confirmation pending legal review."
    )
    return doc

# -------------------------------------------------------------
# 7. NL-MES-OMSOP-001: Operation and Maintenance SOP
# -------------------------------------------------------------
def build_omsop():
    doc = Document()
    doc_id = "NL-MES-OMSOP-001"
    title = "Operation and Maintenance SOP - Novo Life MES PAS-X"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "1.0", "Draft", "2025-01-25")

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "Governs operational procedures for PAS-X MES including user account administration, "
        "hourly database backup execution, monthly disaster recovery drill execution, "
        "and biennial periodic review procedures."
    )
    return doc

# -------------------------------------------------------------
# 8. NL-MES-SLA-001: Service Level Agreement
# -------------------------------------------------------------
def build_sla():
    doc = Document()
    doc_id = "NL-MES-SLA-001"
    title = "Service Level Agreement (SLA) - Novo Life MES IT Operations"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "0.9", "Pre-Operational Draft", "2025-02-15")

    doc.add_heading("1. Operational Status & Scope", level=1)
    p_status = doc.add_paragraph()
    r = p_status.add_run("SYSTEM OPERATIONAL STATUS: PRE-OPERATIONAL / NOT ACTIVATED\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(160, 40, 40)

    doc.add_paragraph(
        "The operational SLA between Global Packaging Operations and IT Infrastructure is currently "
        "in an unactivated state. Operational support tiers (Level 1, 2, and 3) will only be activated "
        "upon successful completion of Gate G6 (Operational Handover) and formal Quality Unit release."
    )
    return doc

# -------------------------------------------------------------
# 9. NL-MES-ITPSE-001: IT Periodic System Evaluation
# -------------------------------------------------------------
def build_itpse():
    doc = Document()
    doc_id = "NL-MES-ITPSE-001"
    title = "IT Periodic System Evaluation (ITPSE) - Pre-Go-Live Readiness Audit"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "0.1", "Evaluation Audit Complete", "2025-03-12")

    doc.add_heading("1. Overall Conclusion & Release Recommendation", level=1)
    p_rec = doc.add_paragraph()
    r = p_rec.add_run("CURRENT RELEASE RECOMMENDATION: HOLD / DEFER - DO NOT RELEASE\n")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(180, 20, 20)

    doc.add_paragraph(
        "Based on comprehensive audit of computerized system qualification evidence, the Novo Life MES PAS-X (SYS-MES-001) "
        "is evaluated as:\n"
        "• NOT IN CONTROL\n"
        "• NOT VALIDATED\n"
        "• NOT FIT FOR OPERATIONAL USE\n"
        "• NOT RELEASED\n\n"
        "The system cannot be released for commercial batch packaging until all open verification activities "
        "are executed and residual risks are formally accepted."
    )

    doc.add_heading("2. Primary Regulatory Findings", level=1)
    doc.add_paragraph(
        "1. Intended-Use Verification Gap: Operational verification (OV / PfV / UAT) has not been performed.\n"
        "2. Lifecycle Gate Failures: Gate G5 (Release Readiness) and Gate G6 (Operational Handover) are marked NOT MET.\n"
        "3. Residual Risk Sign-Off: All 49 working high risks in ITRRA have unrated residual risks and lack Quality Unit sign-off.\n"
        "4. Validation Summary Report: VSR has been deferred."
    )
    return doc

# -------------------------------------------------------------
# 10. NL-MES-IREP-001: IT Implementation Report
# -------------------------------------------------------------
def build_irep():
    doc = Document()
    doc_id = "NL-MES-IREP-001"
    title = "IT Implementation Report - Novo Life MES PAS-X (G5 / G6 Evaluation)"
    apply_dummy_watermark(doc, doc_id, title)
    add_title_block(doc, doc_id, title, "1.0", "Published (Blocked)", "2025-03-14")

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This Implementation Report summarizes the technical deployment, installation qualification, "
        "and phase gate evaluations for Novo Life MES PAS-X (SYS-MES-001)."
    )

    doc.add_heading("2. Technical Verification Status", level=1)
    doc.add_paragraph(
        "• Infrastructure & Network Provisioning: COMPLETE\n"
        "• Installation Qualification (IQ): COMPLETE\n"
        "• ERP (SAP S/4HANA) Technical Integration: COMPLETE\n"
        "• LIMS Technical Integration: COMPLETE\n"
        "• Backup and Disaster Recovery Failover Testing: COMPLETE\n"
        "• Intended-Use Verification (OV / PfV / UAT): NOT PERFORMED (Blocked by open recipe test scripts)\n"
        "• Production Deployment: NOT PERFORMED"
    )

    doc.add_heading("3. Operational Readiness & Handover Items", level=1)
    doc.add_paragraph(
        "• Operational Shopfloor Training: OPEN (0 of 250 operators trained)\n"
        "• Operational Handover to Business: OPEN\n"
        "• Residual Risk Acceptance: OPEN\n"
        "• Validation Summary Report (VSR): DEFERRED"
    )

    doc.add_heading("4. Lifecycle Phase Gate Status", level=1)
    doc.add_paragraph(
        "• Gate G1 (Inception): MET\n"
        "• Gate G2 (Requirements Baseline): MET\n"
        "• Gate G3 (Detailed Design Freeze): MET\n"
        "• Gate G4 (Technical Verification): MET (Technical IQ / Interfaces only)\n"
        "• Gate G5 (Release Readiness): NOT MET — Blocked by missing OV/PfV/UAT and unrated residual risks\n"
        "• Gate G6 (Operational Handover): NOT MET — Blocked by open training and unactivated SLA"
    )
    return doc

def generate_all_mes_documents():
    dest_dir = settings.SAMPLE_DOCS_DIR
    os.makedirs(dest_dir, exist_ok=True)
    print(f"Generating authentic Novo Life MES PAS-X dummy documents in {dest_dir}...")

    docs = [
        ("NL-MES-MLGP-001.docx", build_mlgp()),
        ("NL-MES-URS-001.docx", build_urs()),
        ("NL-MES-FS-001.docx", build_fs()),
        ("NL-MES-ITRA-001.docx", build_itra()),
        ("NL-MES-ITRRA-001.docx", build_itrra()),
        ("NL-MES-SUPA-001.docx", build_supa()),
        ("NL-MES-OMSOP-001.docx", build_omsop()),
        ("NL-MES-SLA-001.docx", build_sla()),
        ("NL-MES-ITPSE-001.docx", build_itpse()),
        ("NL-MES-IREP-001.docx", build_irep())
    ]

    for fname, doc_obj in docs:
        out_path = os.path.join(dest_dir, fname)
        doc_obj.save(out_path)
        print(f"  [+] Created {fname} ({os.path.getsize(out_path)} bytes)")

    print("All 10 authentic Novo Life MES PAS-X dummy documents generated successfully.")

if __name__ == "__main__":
    generate_all_mes_documents()
