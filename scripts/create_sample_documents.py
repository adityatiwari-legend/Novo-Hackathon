"""
Script to generate realistic GxP sample documents for Novo Nordisk Hackathon demo:
1. System_A_URS.docx (contains intentional gap: QA Approval missing/unapproved)
2. System_A_Risk_Assessment.docx
3. SOP_Change_Control.docx
4. SOP_Document_Management.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_system_a_urs(output_path: str):
    doc = Document()
    
    title = doc.add_heading('User Requirements Specification (URS)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('System A: Validated Laboratory Information Management System (LIMS)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Document Ref: URS-SYS-A-001 | Version: 1.0 | Status: Draft / Pending Review')
    
    doc.add_heading('1. System Overview & Scope', level=1)
    doc.add_paragraph(
        'System A is an enterprise Laboratory Information Management System (LIMS) utilized in QC analytical laboratories '
        'to manage sample registration, test execution, instrument integration, calculation verification, and batch release reporting. '
        'This system is categorized as GxP-Critical (Direct GxP Impact) pursuant to GAMP 5 Category 4 configured software.'
    )
    
    doc.add_heading('2. Regulatory Compliance Framework', level=1)
    doc.add_paragraph(
        'System A must comply with 21 CFR Part 11, EU Annex 11, and ALCOA+ data integrity principles (Attributable, Legible, '
        'Contemporaneous, Original, Accurate, Complete, Consistent, Enduring, Available). All analytical data generated must remain immutable '
        'and subject to time-stamped, append-only electronic audit trails.'
    )
    
    doc.add_heading('3. Functional Requirements', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Req ID'
    hdr_cells[1].text = 'Requirement Statement'
    hdr_cells[2].text = 'Criticality'
    for cell in hdr_cells:
        set_cell_background(cell, '003366')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                
    reqs = [
        ('REQ-001', 'System shall enforce unique user credentials and multi-factor authentication (MFA).', 'High'),
        ('REQ-002', 'System shall generate immutable audit trail records for all creation, modification, and deletion of GxP records.', 'Critical'),
        ('REQ-003', 'System shall require dual-authorized electronic signatures for final analytical batch certificate approval.', 'Critical'),
        ('REQ-004', 'System shall support automated integration with chromatography data systems via secure REST endpoints.', 'Medium'),
    ]
    for req_id, text, crit in reqs:
        row_cells = table.add_row().cells
        row_cells[0].text = req_id
        row_cells[1].text = text
        row_cells[2].text = crit
        
    doc.add_heading('4. Data Integrity & Security Controls', level=1)
    doc.add_paragraph(
        'Role-Based Access Control (RBAC) must enforce segregation of duties between System Administrators, '
        'QC Analysts, and QA Reviewers. System Administrators shall not have privileges to alter or suppress laboratory test results.'
    )
    
    doc.add_heading('5. System Ownership & Roles', level=1)
    doc.add_paragraph('Business Owner: Dr. Marcus Vance (Head of Analytical QC)')
    doc.add_paragraph('Technical System Owner: Sarah Jenkins (IT Lead)')
    doc.add_paragraph('QA Compliance Contact: Pending Quality Unit Assignment')
    
    # Intentional Gap in Section 6: QA approval is MISSING
    doc.add_heading('6. Document Approvals & Signatures', level=1)
    doc.add_paragraph(
        'This User Requirements Specification requires formal review and signature sign-off from the Business Owner, '
        'the Technical System Owner, and the QA Compliance Unit prior to operational qualification and formal GxP audit readiness.'
    )
    
    app_table = doc.add_table(rows=1, cols=4)
    app_hdr = app_table.rows[0].cells
    app_hdr[0].text = 'Role'
    app_hdr[1].text = 'Name'
    app_hdr[2].text = 'Signature Status'
    app_hdr[3].text = 'Date'
    for cell in app_hdr:
        set_cell_background(cell, '4682B4')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                
    sig_rows = [
        ('Business Owner', 'Dr. Marcus Vance', 'SIGNED (Approved)', '2026-06-15'),
        ('Technical Owner', 'Sarah Jenkins', 'SIGNED (Approved)', '2026-06-16'),
        ('QA Compliance Unit', 'Not Assigned', 'MISSING - NOT APPROVED', 'Not found'),
    ]
    for role, name, status, dt in sig_rows:
        row = app_table.add_row().cells
        row[0].text = role
        row[1].text = name
        row[2].text = status
        row[3].text = dt
        if 'MISSING' in status:
            set_cell_background(row[2], 'FFCCCC')
            
    doc.add_paragraph('\n[Notice: QA Approval section incomplete. Pending formal QA validation sign-off.]')
    doc.save(output_path)
    print(f"Created {output_path}")

def create_system_a_risk_assessment(output_path: str):
    doc = Document()
    title = doc.add_heading('GxP IT System Risk Assessment', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('System A: Validated LIMS | Document Ref: RA-SYS-A-001 | Version: 1.0')
    doc.add_paragraph('Author: IT Quality Assurance & Validation Team | Status: Approved | Date: 2026-06-20')
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This GxP Risk Assessment evaluates System A under GAMP 5 and ICH Q9 Quality Risk Management principles. '
        'Because System A directly manages product release data and batch disposition analytical records, '
        'system failures or unauthorized modifications pose direct patient safety and regulatory compliance risks.'
    )
    
    doc.add_heading('2. Risk Evaluation Matrix', level=1)
    table = doc.add_table(rows=1, cols=5)
    for i, name in enumerate(['Risk ID', 'Hazard Description', 'Severity', 'Likelihood', 'Mitigation Control']):
        c = table.rows[0].cells[i]
        c.text = name
        set_cell_background(c, '003366')
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
                
    risks = [
        ('RA-01', 'Unapproved URS leading to untested functionality', 'High', 'Medium', 'Mandate QA Approval prior to operational qualification.'),
        ('RA-02', 'Audit trail bypass during batch record calculations', 'Critical', 'Low', 'Hardware write-once logging & SHA-256 database triggers.'),
        ('RA-03', 'Stale SOP documentation causing operator deviation', 'Medium', 'Medium', 'Continuous document lifecycle monitoring & annual reviews.'),
    ]
    for rid, haz, sev, lk, mit in risks:
        r = table.add_row().cells
        r[0].text = rid
        r[1].text = haz
        r[2].text = sev
        r[3].text = lk
        r[4].text = mit
        
    doc.add_heading('3. Control Mappings & ALCOA+ Standards', level=1)
    doc.add_paragraph(
        'Control C-01: All requirements in URS-SYS-A-001 must be linked to verified IQ/OQ/PQ test scripts. '
        'Control C-02: User credentials and segregation of duties must be verified biannually.'
    )
    doc.save(output_path)
    print(f"Created {output_path}")

def create_sop_change_control(output_path: str):
    doc = Document()
    title = doc.add_heading('Standard Operating Procedure: GxP IT Change Control', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Document Ref: SOP-CC-004 | Version: 2.4 | Status: Effective | Effective Date: 2025-03-10')
    doc.add_paragraph('Review Cycle: Biennial | Next Review: 2027-03-10 | Owner: IT Quality Compliance')
    
    doc.add_heading('1. Purpose & Scope', level=1)
    doc.add_paragraph(
        'This procedure establishes mandatory requirements for initiating, evaluating, approving, implementing, '
        'and testing changes to computerized systems utilized in GxP environments.'
    )
    
    doc.add_heading('2. Change Classification & Severity', level=1)
    doc.add_paragraph(
        'All changes to validated software, hardware, or network configurations must be classified as Minor, Major, or Emergency. '
        'Major changes affecting GxP data workflows require full validation regression testing and formal QA pre- and post-approval.'
    )
    
    doc.add_heading('3. Approval Workflow Requirements', level=1)
    doc.add_paragraph(
        'No GxP computerized system change may be migrated to production without documented approval from the System Owner '
        'and Quality Assurance. Automated changes or AI-assisted updates must be gated by human authorization.'
    )
    doc.save(output_path)
    print(f"Created {output_path}")

def create_sop_document_management(output_path: str):
    doc = Document()
    title = doc.add_heading('Standard Operating Procedure: Document Management & Periodic Review', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Document Ref: SOP-DM-001 | Version: 3.2 | Status: Effective | Effective Date: 2024-09-01')
    doc.add_paragraph('Review Cycle: Biennial | Next Review: 2026-09-01 | Owner: Global Quality Assurance')
    
    doc.add_heading('1. Purpose', level=1)
    doc.add_paragraph(
        'Defines lifecycle controls for GxP documents including drafting, versioning, review, QA approval, distribution, and archival.'
    )
    
    doc.add_heading('2. Periodic Review Requirements', level=1)
    doc.add_paragraph(
        'All GxP documentation must undergo periodic review at least once every 24 months. If a document passes its periodic review '
        'due date without a formalized review extension or revised version, its status transitions to Overdue and non-compliant.'
    )
    
    doc.add_heading('3. Signature & Traceability Controls', level=1)
    doc.add_paragraph(
        'Every GxP document must have clear author, reviewer, and QA approver signatures. Missing approval records invalidate document readiness.'
    )
    doc.save(output_path)
    print(f"Created {output_path}")

if __name__ == '__main__':
    target_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'sample_documents')
    os.makedirs(target_dir, exist_ok=True)
    create_system_a_urs(os.path.join(target_dir, 'System_A_URS.docx'))
    create_system_a_risk_assessment(os.path.join(target_dir, 'System_A_Risk_Assessment.docx'))
    create_sop_change_control(os.path.join(target_dir, 'SOP_Change_Control.docx'))
    create_sop_document_management(os.path.join(target_dir, 'SOP_Document_Management.docx'))
    print("All sample GxP documents successfully generated.")
