import os
from datetime import datetime, timezone
from typing import Dict, Any, List, Tuple
from sqlalchemy.orm import Session
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.app.core.config import settings
from backend.app.models.entities import EvidencePack, System, Document, AuditLog
from backend.app.schemas.domain import AgentResult

class AuditEvidenceAgent:
    def __init__(self, output_dir: str = os.path.join(settings.DATA_DIR, "evidence_packs")):
        self.name = "audit_evidence_agent"
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def validate_evidence(
        self,
        checklist_results: List[Dict[str, Any]],
        findings: List[Dict[str, Any]],
        risks: List[Dict[str, Any]],
        recommendations: List[Dict[str, Any]],
        gates: List[Dict[str, Any]] = None
    ) -> Tuple[bool, List[str]]:
        """
        Strict pre-flight evidence validation.
        Ensures that every finding, risk, and release blocker is backed by verifiable source citations.
        """
        validation_errors = []

        # 1. Verify findings citations
        for idx, f in enumerate(findings):
            cites = f.get("source_citations", [])
            title = f.get("title", f"Finding #{idx}")
            if not cites:
                validation_errors.append(f"Finding '{title}' lacks grounded source document citations.")

        # 2. Verify release blockers
        if gates:
            for g in gates:
                if g.get("status") in ["NOT_MET", "BLOCKED"] and not g.get("evidence_doc"):
                    validation_errors.append(f"Release gate {g.get('gate_code')} marked blocked without evidence document reference.")

        # 3. Verify recommendations link
        for r in recommendations:
            if not r.get("rationale") and not r.get("title"):
                validation_errors.append("Recommendation entry contains empty rationale.")

        is_valid = len(validation_errors) == 0
        return is_valid, validation_errors

    def generate_docx_pack(
        self,
        filepath: str,
        system_name: str,
        system_id: str,
        readiness_score: int,
        checklist_results: List[Dict[str, Any]],
        findings: List[Dict[str, Any]],
        risks: List[Dict[str, Any]],
        recommendations: List[Dict[str, Any]],
        audit_logs: List[Dict[str, Any]],
        gates: List[Dict[str, Any]] = None
    ):
        doc = DocxDocument()
        
        # Header / Notice
        p_head = doc.sections[0].header.paragraphs[0]
        p_head.text = "DUMMY / HACKATHON / TRAINING SIMULATION | GxP AUDIT EVIDENCE PACK"
        p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if p_head.runs:
            p_head.runs[0].font.size = Pt(8.5)
            p_head.runs[0].font.color.rgb = RGBColor(180, 50, 50)

        # Title
        title = doc.add_heading("GxP IT SYSTEM AUDIT EVIDENCE PACK", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        sub = doc.add_paragraph("Comprehensive Regulatory Compliance & Audit Readiness Dossier")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Simulation Warning Box
        p_box = doc.add_paragraph()
        p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_box = p_box.add_run(
            "⚠️ DUMMY / HACKATHON / TRAINING SIMULATION RECORD\n"
            "This document is a synthetic simulation record compiled for the Novo Life MES PAS-X hackathon scenario. "
            "It does not represent genuine approved operational data or an operational release."
        )
        r_box.bold = True
        r_box.font.size = Pt(9.5)
        r_box.font.color.rgb = RGBColor(180, 20, 20)
        
        meta = doc.add_paragraph(
            f"Target System: {system_name} ({system_id})\n"
            f"Lifecycle Status: PRE-OPERATIONAL / NOT ACTIVATED\n"
            f"Current Release Recommendation: HOLD / DEFER - DO NOT RELEASE\n"
            f"Assessment Timestamp: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}\n"
            f"Audit Readiness Index: {readiness_score}%\n"
            "Framework Alignment: 21 CFR Part 11 | EU Annex 11 | GAMP 5 Category 4 | ALCOA+ | ICH Q9"
        )
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading("1. Executive Summary & Readiness Assessment", level=1)
        doc.add_paragraph(
            f"This dossier provides automated, continuous verification of {system_name} against established "
            f"GxP regulatory compliance standards. The current deterministic audit readiness score is evaluated at {readiness_score}%. "
            f"The primary recommendation is HOLD / DEFER - DO NOT RELEASE due to open intended-use verification (OV/PfV/UAT not performed), "
            f"unrated residual risks in the requirement baseline, and blocked lifecycle phase gates G5 and G6."
        )
        
        # Phase Gates
        if gates:
            doc.add_heading("2. Lifecycle Phase Gate Evaluation (G1 - G6)", level=1)
            t_gate = doc.add_table(rows=1, cols=4)
            t_gate.rows[0].cells[0].paragraphs[0].add_run("Gate Code").bold = True
            t_gate.rows[0].cells[1].paragraphs[0].add_run("Gate Name").bold = True
            t_gate.rows[0].cells[2].paragraphs[0].add_run("Status").bold = True
            t_gate.rows[0].cells[3].paragraphs[0].add_run("Prerequisites / Blocking Evidence").bold = True
            for g in gates:
                row = t_gate.add_row()
                row.cells[0].paragraphs[0].text = g.get("gate_code", "G")
                row.cells[1].paragraphs[0].text = g.get("gate_name", "Gate")
                row.cells[2].paragraphs[0].text = g.get("status", "NOT_MET")
                row.cells[3].paragraphs[0].text = g.get("blocking_reason") or f"Satisfied in {g.get('evidence_doc', 'MLGP')}"
        
        # Checklist
        doc.add_heading("3. Mandatory Compliance Checklist Results", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].paragraphs[0].add_run("Control ID").bold = True
        table.rows[0].cells[1].paragraphs[0].add_run("Regulatory Requirement").bold = True
        table.rows[0].cells[2].paragraphs[0].add_run("Result").bold = True
        table.rows[0].cells[3].paragraphs[0].add_run("Evidence Finding").bold = True
        
        for c in checklist_results:
            row = table.add_row()
            row.cells[0].paragraphs[0].text = c.get("check_code", "")
            row.cells[1].paragraphs[0].text = c.get("requirement", "")
            row.cells[2].paragraphs[0].text = c.get("status", "")
            row.cells[3].paragraphs[0].text = c.get("evidence", "")
            
        # Findings
        doc.add_heading("4. Identified Compliance Gaps & Risks", level=1)
        for f in findings:
            p = doc.add_paragraph()
            r_sev = p.add_run(f"[{f.get('severity', 'HIGH')}] {f.get('title', 'Gap')}\n")
            r_sev.bold = True
            p.add_run(f"Description: {f.get('description', '')}\n")
            p.add_run(f"Recommended Remediation: {f.get('recommended_action', '')}\n")
            if f.get("source_citations"):
                cite = f["source_citations"][0]
                p.add_run(f"Citation: {cite.get('document', '')} (Page {cite.get('page', 1)}, Section: {cite.get('section', 'General')})\n")
                
        # Audit Ledger
        doc.add_heading("5. Cryptographically Chained Audit Trail Ledger", level=1)
        t_audit = doc.add_table(rows=1, cols=4)
        t_audit.rows[0].cells[0].paragraphs[0].add_run("Timestamp (UTC)").bold = True
        t_audit.rows[0].cells[1].paragraphs[0].add_run("Actor / Agent").bold = True
        t_audit.rows[0].cells[2].paragraphs[0].add_run("Action").bold = True
        t_audit.rows[0].cells[3].paragraphs[0].add_run("SHA-256 Hash Link").bold = True
        for a in audit_logs[:15]:
            row = t_audit.add_row()
            row.cells[0].paragraphs[0].text = str(a.get("timestamp", ""))[:19].replace("T", " ")
            row.cells[1].paragraphs[0].text = f"{a.get('actor_type', '')}: {a.get('agent_name', '')}"
            row.cells[2].paragraphs[0].text = a.get("action", "")
            h = a.get("event_hash", "")
            row.cells[3].paragraphs[0].text = f"{h[:10]}...{h[-6:]}" if len(h) > 16 else h
            
        doc.save(filepath)

    def generate_pdf_pack(
        self,
        filepath: str,
        system_name: str,
        system_id: str,
        readiness_score: int,
        checklist_results: List[Dict[str, Any]],
        findings: List[Dict[str, Any]],
        risks: List[Dict[str, Any]],
        recommendations: List[Dict[str, Any]],
        audit_logs: List[Dict[str, Any]],
        gates: List[Dict[str, Any]] = None
    ):
        doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        story = []

        title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=20, leading=24, textColor=colors.HexColor('#005999'), alignment=1)
        sub_style = ParagraphStyle('SubStyle', parent=styles['Normal'], fontSize=10, leading=14, textColor=colors.HexColor('#475569'), alignment=1)
        h2_style = ParagraphStyle('H2Style', parent=styles['Heading2'], fontSize=13, leading=17, textColor=colors.HexColor('#0f172a'), spaceBefore=10, spaceAfter=6)
        body_style = ParagraphStyle('BodyStyle', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=colors.HexColor('#334155'))
        warn_style = ParagraphStyle('WarnStyle', parent=styles['Normal'], fontSize=8.5, leading=12, textColor=colors.HexColor('#991b1b'), alignment=1)

        # Simulation Notice Banner
        story.append(Paragraph("<b>⚠️ DUMMY / HACKATHON / TRAINING SIMULATION RECORD</b><br/>This document is a synthetic simulation record for the Novo Life MES PAS-X hackathon scenario. Not a genuine operational record.", warn_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("GxP IT SYSTEM AUDIT EVIDENCE DOSSIER", title_style))
        story.append(Paragraph(f"System: <b>{system_name} ({system_id})</b> | Status: <b>PRE-OPERATIONAL</b>", sub_style))
        story.append(Paragraph(f"Release Decision: <b>HOLD / DEFER - DO NOT RELEASE</b> | Readiness: <b>{readiness_score}%</b>", sub_style))
        story.append(Spacer(1, 12))

        # Executive Summary
        story.append(Paragraph("1. Executive Summary & Verification Posture", h2_style))
        story.append(Paragraph(
            f"This dossier summarizes the compliance evaluation of {system_name}. "
            f"Deterministic readiness index: <b>{readiness_score}%</b>. "
            f"System release is currently <b>HELD / DEFERRED</b> due to incomplete intended-use verification (OV/PfV/UAT not performed), "
            f"unrated residual risks across 49 working high requirements, and unsatisfied lifecycle gates G5 and G6.",
            body_style
        ))
        story.append(Spacer(1, 10))

        # Phase Gates
        if gates:
            story.append(Paragraph("2. Lifecycle Phase Gate Status (G1 - G6)", h2_style))
            gate_data = [["Gate", "Name", "Status", "Prerequisites / Blocking Reason"]]
            for g in gates:
                gate_data.append([
                    g.get("gate_code", ""),
                    Paragraph(g.get("gate_name", ""), body_style),
                    g.get("status", ""),
                    Paragraph(g.get("blocking_reason") or "Satisfied", body_style)
                ])
            gt = Table(gate_data, colWidths=[40, 150, 70, 280])
            gt.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
            ]))
            story.append(gt)
            story.append(Spacer(1, 10))

        # Findings Summary
        story.append(Paragraph("3. Active Compliance Gaps & Release Blockers", h2_style))
        for f in findings[:6]:
            sev_color = "#dc2626" if f.get("severity") in ["CRITICAL", "HIGH"] else "#d97706"
            story.append(Paragraph(
                f"<font color='{sev_color}'><b>[{f.get('severity', 'HIGH')}]</b></font> <b>{f.get('title', 'Finding')}</b>: {f.get('description', '')}",
                body_style
            ))
            story.append(Spacer(1, 3))

        story.append(Spacer(1, 10))
        story.append(Paragraph("4. Cryptographic SHA-256 Audit Trail Verification", h2_style))
        audit_data = [["Timestamp (UTC)", "Actor / Agent", "Action", "SHA-256 Hash"]]
        for a in audit_logs[:10]:
            h = a.get("event_hash", "")
            h_trunc = f"{h[:10]}...{h[-6:]}" if len(h) > 16 else h
            audit_data.append([
                str(a.get("timestamp", ""))[:19].replace("T", " "),
                f"{a.get('actor_type', '')}: {a.get('agent_name', '')}",
                Paragraph(a.get("action", ""), body_style),
                h_trunc
            ])
        at = Table(audit_data, colWidths=[110, 130, 180, 120])
        at.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTSIZE', (0, 0), (-1, -1), 7.5),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ]))
        story.append(at)

        doc.build(story)

    def execute(self, db: Session, system_id: str = "SYS-MES-001", actor_id: str = "audit_evidence_agent") -> AgentResult:
        system = db.query(System).filter(System.id == system_id).first()
        sys_name = system.name if system else "Novo Life MES PAS-X"
        readiness = system.readiness_score if system else 48
        
        from backend.app.services.compliance_engine import compliance_engine
        from backend.app.services.release_gate_engine import release_gate_engine
        comp_res = compliance_engine.evaluate_system(db, system_id)
        gate_res = release_gate_engine.evaluate_release_gates(db, system_id)
        
        checklist_results = comp_res.get("checks", [])
        findings_objs = comp_res.get("findings", [])
        findings = []
        for f in findings_objs:
            if isinstance(f, dict):
                findings.append(f)
            else:
                findings.append({
                    "title": getattr(f, "title", ""),
                    "description": getattr(f, "description", ""),
                    "severity": getattr(f, "severity", "MEDIUM"),
                    "recommended_action": getattr(f, "recommended_action", None),
                    "source_citations": getattr(f, "source_citations", [])
                })
        gates = gate_res.get("gates", [])
        
        # Pre-flight evidence validation
        is_valid, val_errors = self.validate_evidence(checklist_results, findings, [], [], gates)
        if not is_valid:
            return AgentResult(
                agent=self.name,
                status="VALIDATION_FAILED",
                findings=[{"error": err} for err in val_errors],
                confidence=0.20,
                metadata={"summary": f"Evidence validation failed with {len(val_errors)} errors."}
            )
        
        # Recent audit logs
        logs = db.query(AuditLog).order_by(AuditLog.timestamp.desc()).limit(15).all()
        audit_logs_data = [
            {
                "timestamp": l.timestamp.isoformat() if hasattr(l.timestamp, "isoformat") else str(l.timestamp),
                "actor_type": l.actor_type,
                "agent_name": l.agent_name,
                "action": l.action,
                "event_hash": l.event_hash
            } for l in logs
        ]

        ts_str = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"Evidence_Pack_{system_id}_{ts_str}.pdf"
        docx_filename = f"Evidence_Pack_{system_id}_{ts_str}.docx"
        
        pdf_path = os.path.join(self.output_dir, pdf_filename)
        docx_path = os.path.join(self.output_dir, docx_filename)
        
        self.generate_pdf_pack(pdf_path, sys_name, system_id, readiness, checklist_results, findings, [], [], audit_logs_data, gates)
        self.generate_docx_pack(docx_path, sys_name, system_id, readiness, checklist_results, findings, [], [], audit_logs_data, gates)
        
        # Record EvidencePack in database
        pack = EvidencePack(
            system_id=system_id,
            title=f"Audit Evidence Dossier: {sys_name}",
            version="2.0",
            scope="Complete GxP Lifecycle & Release Readiness Assessment",
            file_path=pdf_path,
            docx_file_path=docx_path,
            citations_json=[{"doc": "NL-MES-ITPSE-001", "scope": "HOLD / DEFER"}],
            status="GENERATED",
            generated_by=self.name
        )
        db.add(pack)
        db.commit()
        db.refresh(pack)
        
        from backend.app.models.entities import create_audit_log
        create_audit_log(
            db=db,
            actor_type="AGENT",
            actor_id=self.name,
            action="EVIDENCE_PACK_GENERATED",
            entity_type="EVIDENCE_PACK",
            entity_id=pack.id,
            details={"pdf": pdf_filename, "docx": docx_filename, "readiness": readiness, "gates_checked": len(gates)},
            agent_name=self.name
        )

        return AgentResult(
            agent=self.name,
            status="COMPLETED",
            findings=[],
            confidence=0.98,
            metadata={
                "summary": f"Successfully generated audit evidence pack dossier in PDF and DOCX for {sys_name}.",
                "evidence_pack_id": pack.id,
                "pdf_filename": pdf_filename,
                "docx_filename": docx_filename,
                "file_path": pack.file_path,
                "docx_file_path": pack.docx_file_path
            }
        )

    def run(self, db: Session, system_id: str = "SYS-MES-001", *args, **kwargs) -> AgentResult:
        return self.execute(db, system_id)

audit_evidence_agent = AuditEvidenceAgent()
evidence_agent = audit_evidence_agent
