"""
Audit Report Generation Service for GxP IT Systems.
Generates comprehensive 14-section PDF and DOCX Audit Reports
incorporating dynamic audit readiness scores, question-by-question evidence,
risk summaries, lifecycle deviations, and regulatory disclaimers.
"""

import os
import logging
from datetime import datetime, timezone
from typing import Dict, Any, List, Optional
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.app.core.config import settings
from backend.app.services.audit_engine import audit_engine
from backend.app.schemas.domain import AuditAssessmentResponse, CrossDocComparisonResponse

DISCLAIMER_TEXT = (
    "This is a hackathon/training simulation and does not constitute a regulatory audit, "
    "validation decision, or production release authorization."
)

class AuditReportService:
    def __init__(self, output_dir: str = os.path.join(settings.DATA_DIR, "evidence_packs")):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_pdf_report(
        self,
        filepath: str,
        assessment: AuditAssessmentResponse,
        comparison: CrossDocComparisonResponse
    ):
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=36,
            leftMargin=36,
            topMargin=36,
            bottomMargin=36
        )
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            textColor=colors.HexColor("#002B49"),
            alignment=1,
            spaceAfter=6
        )
        subtitle_style = ParagraphStyle(
            'ReportSub',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#475569"),
            alignment=1,
            spaceAfter=12
        )
        sec_heading = ParagraphStyle(
            'SecHeading',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#002B49"),
            spaceBefore=12,
            spaceAfter=6
        )
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=8.5,
            leading=11.5,
            textColor=colors.HexColor("#1E293B")
        )
        body_bold = ParagraphStyle(
            'BodyBold',
            parent=body_style,
            fontName='Helvetica-Bold'
        )
        disclaimer_style = ParagraphStyle(
            'Disclaimer',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#B91C1C"),
            alignment=1
        )

        elements = []

        # Header Disclaimer
        elements.append(Paragraph(f"⚠️ DUMMY / HACKATHON SIMULATION RECORD: {DISCLAIMER_TEXT}", disclaimer_style))
        elements.append(Spacer(1, 8))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CBD5E1"), spaceAfter=10))

        # Title Block
        elements.append(Paragraph("GxP IT AUDIT & LIFECYCLE INTELLIGENCE REPORT", title_style))
        elements.append(Paragraph(f"System: Novo Life MES PAS-X (SYS-MES-001) | Checklist: Top 25 Difficult-Auditor Questions", subtitle_style))
        elements.append(Spacer(1, 10))

        # 1. Executive Summary & KPIs
        elements.append(Paragraph("1. Executive Summary", sec_heading))
        summary_text = (
            f"An independent audit evaluation was conducted on <b>Novo Life MES PAS-X (SYS-MES-001)</b> using the "
            f"<b>Top 25 Difficult-Auditor GxP IT Checklist (2026)</b>, benchmarked against <b>NN Master IT System Lifecycle SOP "
            f"(HACK-IT-SOP-001)</b> and the <b>GxP LIMS Lifecycle Documentation Package (LIMS-LCP-001)</b>. "
            f"The calculated <b>Overall Audit Readiness Score is {assessment.readiness_score}%</b>. "
            f"Due to critical incomplete verification activities and unrated residual risks, the formal system release "
            f"recommendation remains: <b>HOLD / DEFER - DO NOT RELEASE TO PRODUCTION</b>."
        )
        elements.append(Paragraph(summary_text, body_style))
        elements.append(Spacer(1, 8))

        kpi_data = [
            ["Readiness Score", "Total Questions", "Passed", "Partial", "Failed", "Critical Findings", "Release Decision"],
            [
                f"{assessment.readiness_score}%",
                str(assessment.total_questions),
                str(assessment.passed_count),
                str(assessment.partial_count),
                str(assessment.failed_count),
                str(assessment.critical_findings_count),
                "HOLD / DEFER"
            ]
        ]
        kpi_table = Table(kpi_data, colWidths=[75, 75, 55, 55, 55, 90, 135])
        kpi_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#002B49")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('BACKGROUND', (0, 1), (-1, 1), colors.HexColor("#F8FAFC")),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
            ('TEXTCOLOR', (0, 1), (0, 1), colors.HexColor("#1D4ED8")),
            ('TEXTCOLOR', (5, 1), (6, 1), colors.HexColor("#B91C1C")),
        ]))
        elements.append(kpi_table)
        elements.append(Spacer(1, 12))

        # 2. System Assessed & 3. Assessment Date
        elements.append(Paragraph("2. System Assessed & 3. Assessment Metadata", sec_heading))
        meta_data = [
            ["System ID / Name", "SYS-MES-001 | Novo Life MES PAS-X (Werum PAS-X Category 4)"],
            ["Current Status", "PRE-OPERATIONAL / NOT ACTIVATED [NL-MES-SLA-001 p.1]"],
            ["Assessment Date", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")],
            ["Auditing Engine", "GxP IT Audit & Lifecycle Intelligence Assistant (Deterministic Engine)"],
            ["Predicate Rules", "21 CFR Part 11, 21 CFR 211.68, EU GMP Annex 11 / 15, GAMP 5, ICH Q9(R1)"]
        ]
        meta_table = Table(meta_data, colWidths=[150, 390])
        meta_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor("#F1F5F9")),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        elements.append(meta_table)
        elements.append(Spacer(1, 12))

        # 4. Documents Reviewed
        elements.append(Paragraph("4. Documents Reviewed Across Knowledge Hierarchy", sec_heading))
        docs_data = [
            ["Scope", "Document ID", "Title", "Version", "Role / Status"],
            ["Primary Evidence", "NL-MES-URS-001", "User Requirements Specification", "v1.0", "Draft / Missing QA Approval"],
            ["Primary Evidence", "NL-MES-MLGP-001", "Master Lifecycle Generation Plan", "v1.0", "Approved (Simulated)"],
            ["Primary Evidence", "NL-MES-ITRA-001", "IT Risk Assessment Baseline", "v1.0", "Approved (Simulated)"],
            ["Primary Evidence", "NL-MES-ITRRA-001", "Requirement Risk Assessment", "v1.0", "Residual Risk NOT RATED"],
            ["Primary Evidence", "NL-MES-IREP-001", "IT Implementation Report", "v1.0", "Gate G5 Blocked / OV Open"],
            ["Primary Evidence", "NL-MES-ITPSE-001", "Periodic System Evaluation", "v1.0", "HOLD / DEFER Conclusion"],
            ["Primary Evidence", "NL-MES-SLA-001", "Service Level Agreement", "v0.9", "Pre-Operational / Unactivated"],
            ["Governance SOP", "HACK-IT-SOP-001", "Manage IT System Lifecycle SOP", "v0.1", "Master Governance Standard"],
            ["Benchmark Ref", "LIMS-LCP-001", "GxP LIMS Lifecycle Package", "v0.1", "Laboratory Benchmark Ref"],
            ["Audit Checklist", "CKL-TOP25-2026", "Top 25 Difficult-Auditor Questions", "2026.1", "Master Audit Instrument"]
        ]
        docs_table = Table(docs_data, colWidths=[90, 95, 175, 45, 135])
        docs_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#002B49")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 7.5),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
        ]))
        elements.append(docs_table)
        elements.append(Spacer(1, 14))

        # 5 & 6. Question-by-Question Results
        elements.append(Paragraph("5 & 6. Audit Checklist: Question-by-Question Results", sec_heading))
        q_table_data = [["#", "Q_ID", "Audit Domain / Topic", "Result", "Evidence Citation", "Severity"]]
        for item in assessment.items:
            res_color = colors.HexColor("#166534") if item.status == "PASS" else colors.HexColor("#991B1B") if item.status == "FAIL" else colors.HexColor("#854D0E")
            cite = item.evidence_citations[0] if item.evidence_citations else "Not evidenced"
            q_table_data.append([
                str(item.sequence),
                item.question_id,
                Paragraph(f"<b>{item.control_topic}</b><br/>{item.audit_question[:75]}...", body_style),
                item.status,
                Paragraph(cite, body_style),
                item.risk_level
            ])

        q_table = Table(q_table_data, colWidths=[20, 50, 200, 55, 160, 55])
        q_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#002B49")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 7),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(q_table)
        elements.append(Spacer(1, 14))

        # 7 & 8. Critical Findings & Evidence Details
        elements.append(Paragraph("7 & 8. Critical Audit Findings & Evidence", sec_heading))
        for f in assessment.findings[:4]:
            f_box = [
                [Paragraph(f"<b>FINDING [{f['question_id']}]: {f['title']}</b>", body_bold), Paragraph(f"Severity: {f['severity']}", body_bold)],
                [Paragraph(f"<b>Gap:</b> {f['gap']}<br/><b>Recommendation:</b> {f['recommendation']}<br/><b>Evidence:</b> {', '.join(f['citations'])}", body_style), ""]
            ]
            t_box = Table(f_box, colWidths=[420, 120])
            t_box.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#FEF2F2")),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor("#991B1B")),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#FCA5A5")),
                ('SPAN', (0, 1), (1, 1)),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            elements.append(t_box)
            elements.append(Spacer(1, 6))

        # 9. Risk Summary & 10. Lifecycle Gaps (Master SOP Comparison)
        elements.append(Paragraph("9. Cross-Document Comparison & Lifecycle Deviations", sec_heading))
        elements.append(Paragraph("Direct comparison of MES PAS-X evidence against Master IT SOP (HACK-IT-SOP-001):", body_style))
        elements.append(Spacer(1, 6))

        comp_data = [["Topic", "SOP Expectation", "MES PAS-X Observed", "Status", "Recommended Action"]]
        for c in comparison.items[:4]:
            comp_data.append([
                Paragraph(f"<b>{c.topic}</b>", body_style),
                Paragraph(c.sop_requirement[:110] + "...", body_style),
                Paragraph(c.mes_observed[:110] + "...", body_style),
                c.alignment_status.replace("_", " "),
                Paragraph(c.recommended_action[:95] + "...", body_style)
            ])
        comp_table = Table(comp_data, colWidths=[95, 120, 120, 85, 120])
        comp_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#002B49")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 7),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(comp_table)
        elements.append(Spacer(1, 14))

        # 11 & 12. Recommendations & Overall Readiness
        elements.append(Paragraph("11 & 12. Corrective Actions Required Prior to Production Release", sec_heading))
        recs = [
            "1. <b>Execute Intended-Use Verification (OV / PfV / UAT):</b> Complete qualified test scripts on packaging lines to unblock Gate G5 [NL-MES-IREP-001 p.2].",
            "2. <b>Conduct Authorized Residual Risk Review:</b> Review and accept the 49 working high risks with the Quality Unit [NL-MES-ITRRA-001 p.3].",
            "3. <b>Finalize and Authorize Validation Summary Report (VSR):</b> Route the VSR for formal QA sign-off prior to Gate G5 authorization.",
            "4. <b>Complete Shopfloor Operator Training & SLA Activation:</b> Train 250 packaging line operators and activate 24/7 SLA operational support [NL-MES-SLA-001 p.1]."
        ]
        for r in recs:
            elements.append(Paragraph(r, body_style))
            elements.append(Spacer(1, 3))

        elements.append(Spacer(1, 8))
        # 13 & 14. Limitations & Source References
        elements.append(Paragraph("13 & 14. Limitations & Regulatory Disclaimer", sec_heading))
        elements.append(Paragraph(
            f"<b>LIMITATIONS:</b> {DISCLAIMER_TEXT} "
            "All findings, scores, and release recommendations are deterministic outputs generated "
            "from synthetic documentation for hackathon demonstration purposes.", body_style
        ))

        doc.build(elements)
        logger.info(f"Generated Audit PDF: {filepath}")

    def generate_docx_report(
        self,
        filepath: str,
        assessment: AuditAssessmentResponse,
        comparison: CrossDocComparisonResponse
    ):
        doc = DocxDocument()

        # Header disclaimer
        p_head = doc.sections[0].header.paragraphs[0]
        p_head.text = f"DUMMY / HACKATHON SIMULATION RECORD: {DISCLAIMER_TEXT}"
        p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if p_head.runs:
            p_head.runs[0].font.size = Pt(8.5)
            p_head.runs[0].font.color.rgb = RGBColor(180, 50, 50)

        # Title
        t = doc.add_heading("GxP IT AUDIT & LIFECYCLE INTELLIGENCE REPORT", level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph("Novo Life MES PAS-X (SYS-MES-001) | Top 25 Audit Checklist Assessment")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Disclaimer Callout
        p_warn = doc.add_paragraph()
        p_warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_warn = p_warn.add_run(f"⚠️ {DISCLAIMER_TEXT}")
        r_warn.bold = True
        r_warn.font.size = Pt(9.5)
        r_warn.font.color.rgb = RGBColor(180, 20, 20)

        # 1. Executive Summary
        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph(
            f"System Evaluated: Novo Life MES PAS-X (SYS-MES-001)\n"
            f"Overall Audit Readiness Score: {assessment.readiness_score}%\n"
            f"Total Questions Evaluated: {assessment.total_questions}\n"
            f"Passed: {assessment.passed_count} | Partial: {assessment.partial_count} | Failed: {assessment.failed_count}\n"
            f"Critical Findings: {assessment.critical_findings_count} | High: {assessment.high_findings_count}\n"
            f"System Release Recommendation: HOLD / DEFER - DO NOT RELEASE TO PRODUCTION"
        )

        # 2 & 3. Metadata
        doc.add_heading("2. System Assessed & 3. Assessment Date", level=1)
        t_meta = doc.add_table(rows=4, cols=2)
        meta_rows = [
            ("System", "Novo Life MES PAS-X (Werum PAS-X GAMP Category 4)"),
            ("Assessment Date", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")),
            ("Governing SOP", "NN Master IT System Lifecycle SOP (HACK-IT-SOP-001)"),
            ("Audit Checklist", "Top 25 Checklists GxP IT Audit Questions 2026")
        ]
        for idx, (k, v) in enumerate(meta_rows):
            t_meta.cell(idx, 0).text = k
            t_meta.cell(idx, 1).text = v

        # 4. Documents Reviewed
        doc.add_heading("4. Documents Reviewed", level=1)
        doc.add_paragraph(
            "- Primary Evidence: NL-MES-URS-001, NL-MES-MLGP-001, NL-MES-ITRA-001, NL-MES-ITRRA-001, NL-MES-IREP-001, NL-MES-ITPSE-001, NL-MES-SLA-001\n"
            "- Governance Standard: NN Master IT System Lifecycle SOP (HACK-IT-SOP-001)\n"
            "- Reference Benchmark: GxP LIMS Lifecycle Documentation Package (LIMS-LCP-001)\n"
            "- Audit Instrument: Top 25 Checklists GxP IT Audit Questions 2026 (Top_25_Checklists_GxP_IT_Audit_Questions_2026.xlsx)"
        )

        # 5 & 6. Checklist Results
        doc.add_heading("5 & 6. Question-by-Question Results", level=1)
        t_q = doc.add_table(rows=len(assessment.items) + 1, cols=5)
        headers = ["#", "Q_ID", "Topic / Question", "Result", "Severity"]
        for c, h in enumerate(headers):
            t_q.cell(0, c).text = h

        for idx, it in enumerate(assessment.items):
            r = idx + 1
            t_q.cell(r, 0).text = str(it.sequence)
            t_q.cell(r, 1).text = it.question_id
            t_q.cell(r, 2).text = f"{it.control_topic}: {it.audit_question[:60]}..."
            t_q.cell(r, 3).text = it.status
            t_q.cell(r, 4).text = it.risk_level

        # 7 & 8. Critical Findings
        doc.add_heading("7 & 8. Critical Findings & Evidence", level=1)
        for f in assessment.findings:
            doc.add_paragraph(
                f"• [{f['severity']}] {f['title']}\n"
                f"  Gap: {f['gap']}\n"
                f"  Recommendation: {f['recommendation']}\n"
                f"  Citations: {', '.join(f['citations'])}"
            )

        # 9 & 10. Cross-Document Comparison
        doc.add_heading("9 & 10. Cross-Document Comparison (Master SOP vs MES)", level=1)
        for c in comparison.items:
            doc.add_paragraph(
                f"• {c.topic} ({c.alignment_status}):\n"
                f"  - SOP Expectation: {c.sop_requirement}\n"
                f"  - Observed State: {c.mes_observed}\n"
                f"  - Recommended Action: {c.recommended_action}"
            )

        # 11 & 12. Recommendations
        doc.add_heading("11 & 12. Corrective Action Plan", level=1)
        doc.add_paragraph(
            "1. Complete Intended-Use Verification (OV / PfV / UAT) on commercial packaging line [NL-MES-IREP-001 Section 3.2].\n"
            "2. Complete Authorized Residual Risk Sign-off for 49 working high requirements with Quality Unit [NL-MES-ITRRA-001 Section 3].\n"
            "3. Obtain Quality Unit VSR authorization to satisfy Gate G5 requirements [NL-MES-IREP-001 Section 4.3].\n"
            "4. Deliver packaging operator training to 250 personnel and activate operational SLA [NL-MES-SLA-001 Section 1]."
        )

        # 13 & 14. Limitations & Disclaimer
        doc.add_heading("13 & 14. Limitations & Regulatory Disclaimer", level=1)
        doc.add_paragraph(DISCLAIMER_TEXT)

        doc.save(filepath)
        logger.info(f"Generated Audit DOCX: {filepath}")

    def generate_full_dossier(
        self,
        db: Session,
        system_id: str = "SYS-MES-001",
        checklist_id: str = "CKL-TOP25-CORE"
    ) -> Dict[str, Any]:
        assessment = audit_engine.get_latest_assessment(db, system_id=system_id)
        comparison = audit_engine.cross_document_comparison(db, system_id=system_id)

        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        pdf_name = f"GxP_IT_Audit_Report_{system_id}_{timestamp}.pdf"
        docx_name = f"GxP_IT_Audit_Report_{system_id}_{timestamp}.docx"

        pdf_path = os.path.join(self.output_dir, pdf_name)
        docx_path = os.path.join(self.output_dir, docx_name)

        self.generate_pdf_report(pdf_path, assessment, comparison)
        self.generate_docx_report(docx_path, assessment, comparison)

        return {
            "success": True,
            "pdf_path": pdf_path,
            "docx_path": docx_path,
            "readiness_score": assessment.readiness_score,
            "summary": f"Audit Report compiled for {system_id}. Score: {assessment.readiness_score}% (HOLD / DEFER).",
            "disclaimer": DISCLAIMER_TEXT
        }

audit_report_service = AuditReportService()
