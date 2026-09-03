import os
from typing import Dict, Any
from datetime import datetime, timezone
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from backend.app.core.config import settings

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

class DraftService:
    def __init__(self, drafts_dir: str = os.path.join(settings.DATA_DIR, "drafts")):
        self.drafts_dir = drafts_dir
        os.makedirs(self.drafts_dir, exist_ok=True)

    def generate_draft_section(
        self,
        section_name: str = "6. QA Approval & Validation Sign-Off",
        document_title: str = "System_A_URS.docx",
        system_name: str = "System A: Validated LIMS",
        finding_context: str = "QA approval missing from URS."
    ) -> Dict[str, Any]:
        """
        Generates an AI draft for the missing section with prominent regulatory disclaimers.
        """
        content = (
            "================================================================================\n"
            "AI-GENERATED DRAFT — NOT APPROVED — REQUIRES HUMAN REVIEW AND AUTHORIZATION\n"
            "Regulation Alignment: 21 CFR Part 11 | EU Annex 11 | GAMP 5 Category 4\n"
            "================================================================================\n\n"
            "SECTION 6: QUALITY ASSURANCE (QA) FORMAL APPROVAL AND VALIDATION SIGN-OFF\n\n"
            "6.1 Quality Unit Assessment & Statement of Compliance:\n"
            f"The Quality Assurance Unit has reviewed this User Requirements Specification for {system_name}. "
            "The functional requirements (REQ-001 through REQ-004), data integrity controls, and electronic signature "
            "mechanisms comply with corporate computerized system validation policy and ALCOA+ principles.\n\n"
            "6.2 Operational Qualification Gate Condition:\n"
            "Execution of Operational Qualification (OQ) test scripts is strictly contingent upon formal electronic "
            "sign-off by the Quality Assurance Unit. No deviations or unapproved changes shall be introduced without "
            "an authorized Change Request pursuant to SOP-CC-004.\n\n"
            "6.3 Required Signature Table:\n"
            "------------------------------------------------------------------------------------------------------\n"
            "Role                      Approver Name            Digital Signature ID       Date (UTC)     Disposition\n"
            "------------------------------------------------------------------------------------------------------\n"
            "Technical System Owner    Sarah Jenkins            SIG-2026-SJ-8891           2026-06-16     APPROVED\n"
            "Business Process Owner    Dr. Marcus Vance         SIG-2026-MV-1042           2026-06-15     APPROVED\n"
            "Quality Assurance Lead    [Pending QA Assignment]  [PENDING SIGNATURE]        [Pending]      PENDING REVIEW\n"
            "------------------------------------------------------------------------------------------------------\n"
        )
        
        return {
            "section_name": section_name,
            "document_title": document_title,
            "system_name": system_name,
            "draft_text": content,
            "status": "DRAFT_PENDING_HUMAN_REVIEW",
            "watermark": "AI GENERATED DRAFT - NOT APPROVED - REQUIRES HUMAN REVIEW",
            "created_at": datetime.now(timezone.utc).isoformat()
        }

    def export_draft_docx(self, draft_text: str, section_name: str, document_title: str) -> str:
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_name = section_name.replace(" ", "_").replace(".", "_")
        filename = f"Draft_{clean_name}_{timestamp_str}.docx"
        filepath = os.path.join(self.drafts_dir, filename)
        
        doc = DocxDocument()
        
        # Prominent Watermark Header
        hdr = doc.add_paragraph()
        r1 = hdr.add_run("⚠️ AI GENERATED DRAFT — NOT APPROVED — REQUIRES HUMAN REVIEW\n")
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(220, 38, 38)
        r1.font.size = Pt(13)
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Target Document: {document_title} | Section: {section_name}")
        doc.add_paragraph(f"Generated at: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
        
        doc.add_heading(section_name, level=1)
        
        for line in draft_text.split("\n"):
            line_str = line.strip()
            if not line_str or line_str.startswith("====") or line_str.startswith("----"):
                continue
            p = doc.add_paragraph(line_str)
            
        doc.save(filepath)
        return filepath

draft_service = DraftService()
